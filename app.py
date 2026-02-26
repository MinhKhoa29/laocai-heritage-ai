# python -m streamlit run app.py
  
import streamlit as st
import json
import base64
import pandas as pd
import os
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from gtts import gTTS
import glob



st.markdown("""
<style>

/* ===== TITLE AI EFFECT ===== */
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2 {

    font-weight: 800 !important;
    letter-spacing: 1.5px;
    color: #ffffff !important;

    animation: glowPulse 2.5s ease-in-out infinite;
}

/* Hiệu ứng tăng giảm ánh sáng */
@keyframes glowPulse {
    0% {
        text-shadow:
            0 0 5px #00f5ff,
            0 0 10px #00f5ff,
            0 0 20px #00f5ff;
    }

    50% {
        text-shadow:
            0 0 10px #00f5ff,
            0 0 25px #00f5ff,
            0 0 40px #00f5ff;
    }

    100% {
        text-shadow:
            0 0 5px #00f5ff,
            0 0 10px #00f5ff,
            0 0 20px #00f5ff;
    }
}

</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>

/* ===== SELECTBOX BOX ===== */
section[data-testid="stSidebar"] .stSelectbox > div > div {
    background: linear-gradient(145deg, rgba(255,255,255,0.08), rgba(0,255,255,0.05)) !important;
    border-radius: 16px !important;
    border: 1px solid rgba(0,255,255,0.4) !important;

    padding: 10px 14px !important;   /* tăng padding */
    min-height: 48px !important;     /* đảm bảo đủ cao */
    display: flex !important;
    align-items: center !important;  /* căn giữa theo chiều dọc */

    transition: all 0.3s ease;
}

/* Text bên trong */
section[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] span {
    color: white !important;
    font-weight: 600;
    font-size: 15px !important;
    line-height: 1.4 !important;  /* tránh bị cắt chữ */
}

/* Hover */
section[data-testid="stSidebar"] .stSelectbox > div > div:hover {
    box-shadow: 0 0 18px rgba(0,255,255,0.6);
    border: 1px solid #00f5ff !important;
}

/* Focus */
section[data-testid="stSidebar"] .stSelectbox > div > div:focus-within {
    box-shadow: 0 0 25px rgba(0,255,255,0.9);
}

/* Dropdown panel */
div[data-baseweb="popover"] {
    background: #04293a !important;
    border-radius: 14px !important;
}

</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>

/* ===== SELECTBOX TEXT GLOW NHẸ ===== */
section[data-testid="stSidebar"] 
div[data-baseweb="select"] > div {

    color: #e0ffff !important;
    font-weight: 600 !important;

    text-shadow:
        0 0 4px rgba(0,255,255,0.5),
        0 0 8px rgba(0,255,255,0.3);

    transition: all 0.3s ease;
}

/* Hover tăng nhẹ ánh sáng */
section[data-testid="stSidebar"] 
div[data-baseweb="select"]:hover > div {

    text-shadow:
        0 0 6px rgba(0,255,255,0.7),
        0 0 12px rgba(0,255,255,0.4);
}

</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>

/* ===== BACKGROUND FUTURE AI ===== */
.stApp {
    background: linear-gradient(135deg, #000814, #001d3d, #003566);
    overflow-x: hidden;
}

/* Hạt chuyển động */
.stApp::before {
    content: "";
    position: fixed;
    width: 200%;
    height: 200%;
    background-image: radial-gradient(white 1px, transparent 1px);
    background-size: 50px 50px;
    opacity: 0.05;
    animation: moveParticles 60s linear infinite;
    top: -50%;
    left: -50%;
}

@keyframes moveParticles {
    from { transform: translate(0,0); }
    to { transform: translate(200px,200px); }
}

/* ===== HEADER GRADIENT ANIMATION ===== */
.ai-header {
    font-size: 50px;
    font-weight: 800;
    text-align: center;
    background: linear-gradient(90deg, #00f5ff, #00ffcc, #00f5ff);
    background-size: 300% 300%;
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    animation: gradientMove 6s ease infinite;
}

@keyframes gradientMove {
    0% { background-position: 0% 50%; }
    50% { background-position: 100% 50%; }
    100% { background-position: 0% 50%; }
}

/* ===== GLASS CARD ===== */
.glass-card {
    background: rgba(255,255,255,0.05);
    backdrop-filter: blur(20px);
    padding: 30px;
    border-radius: 25px;
    border: 1px solid rgba(0,255,255,0.3);
    box-shadow: 0 0 30px rgba(0,255,255,0.2);
    color: white;
    transition: 0.3s;
}

.glass-card:hover {
    transform: translateY(-5px);
    box-shadow: 0 0 40px rgba(0,255,255,0.6);
}

/* ===== LOGO GLOW ===== */
.ai-logo {
    font-size: 28px;
    font-weight: bold;
    color: #00f5ff;
    text-shadow: 0 0 10px #00f5ff,
                 0 0 20px #00f5ff,
                 0 0 40px #00f5ff;
}

</style>
""", unsafe_allow_html=True)


st.markdown("""
<style>

.section-title {
    color: #c084fc;
    border-bottom: 3px solid #a855f7;
    display: inline-block;
    padding-bottom: 5px;
}

.info-card {
    background: linear-gradient(145deg, #1e1b4b, #312e81);
    border-radius: 20px;
    padding: 25px;
    color: white;
    box-shadow: 0 0 30px rgba(168,85,247,0.5);
    border: 1px solid rgba(168,85,247,0.3);
}

</style>
""", unsafe_allow_html=True)



st.markdown("""
    <style>
    /* Giới hạn chiều rộng nội dung */
    .block-container {
        max-width: 1100px;
        margin: auto;
    }
    </style>
    """, unsafe_allow_html=True)


# =========================
# CONFIG
# =========================
st.set_page_config(
    page_title="LAO CAI HERITAGE AI",
    page_icon="logo.png",
    layout="wide"
)


# =========================
# STYLE (FINAL CLEAN VERSION)
# =========================

st.markdown("""
<div style="
    background: linear-gradient(135deg,#00587a,#00b4d8);
    padding:30px;
    border-radius:50px;
    color:white;
    text-align:center;
    margin-bottom:20px;
">
    <h2 style="margin-bottom:1px;">Lào Cai – Vùng Đất Hội Tụ Tinh Hoa Văn Hóa Tây Bắc</h2>
    <p style="font-size:20px; opacity:0.5;">
        Văn hóa • Lịch sử • Di tích • Trải nghiệm
    </p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<style>

/* ===== IMPORT FONT ===== */
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;800&family=Montserrat:wght@400;600;700&display=swap');

html, body, [class*="css"] {
    font-family: 'Montserrat', sans-serif;
}

/* ===== BACKGROUND TRẮNG + LOANG XANH ===== */
.stApp {
    background-color: #ffffff;
    background-image:
        radial-gradient(circle at 30% 40%, rgba(0,136,145,0.18) 0%, transparent 45%),
        radial-gradient(circle at 85% 75%, rgba(0,88,122,0.15) 0%, transparent 50%);
    background-attachment: fixed;
}

/* ===== KHUNG CHÍNH ===== */
.main .block-container {
    background: rgba(255,255,255,0.96);
    padding: 2.5rem;
    border-radius: 20px;
    box-shadow: 0 20px 40px rgba(0,0,0,0.15);
}

/* ===== TITLE ===== */
.big-title {
    font-family: 'Playfair Display', serif;
    font-size: 48px;
    font-weight: 800;
    text-align: center;
    color: #0f3057;
    margin-bottom: 10px;
}

/* ===== SIDEBAR ===== */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0f3057, #00587a);
    padding-top: 20px;
}

/* Sidebar title */
section[data-testid="stSidebar"] h1 {
    font-size: 24px !important;
    font-weight: 800 !important;
    color: white !important;
    margin-bottom: 20px;
}

/* Label */
section[data-testid="stSidebar"] label {
    font-size: 13px !important;
    font-weight: 700 !important;
    color: #b3e5fc !important;
    text-transform: uppercase;
}

/* Radio item */
section[data-testid="stSidebar"] div[role="radiogroup"] label[data-baseweb="radio"] {
    padding: 10px 14px;
    border-radius: 10px;
    margin-bottom: 8px;
    font-size: 15px !important;
    font-weight: 600 !important;
    color: white !important;
    transition: 0.25s ease;
}

/* Hover */
section[data-testid="stSidebar"] div[role="radiogroup"] label[data-baseweb="radio"]:hover {
    background: rgba(255,255,255,0.12);
}

/* Selected */
section[data-testid="stSidebar"] div[role="radiogroup"] label[data-baseweb="radio"][aria-checked="true"] {
    background: rgba(255,255,255,0.18);
    box-shadow: inset 4px 0 0 #ffd700;
    font-weight: 700 !important;
}

/* Radio dot */
section[data-testid="stSidebar"] input[type="radio"] + div {
    background-color: white !important;
}

/* ===== CARD ===== */
.card {
    background: #ffffff;
    border-radius: 18px;
    padding: 22px;
    margin-bottom: 20px;
    box-shadow: 0 12px 30px rgba(0,0,0,0.1);
    transition: 0.3s ease;
}

.card:hover {
    transform: translateY(-6px);
    box-shadow: 0 18px 35px rgba(0,0,0,0.18);
}

/* ===== SECTION TITLE ===== */
.section-title {
    font-size: 20px;
    font-weight: 800;
    color: #00587a;
    margin-top: 20px;
    margin-bottom: 8px;
}

.section-title::after {
    content: "";
    display: block;
    width: 50px;
    height: 3px;
    background: linear-gradient(90deg, #00587a, #00b4d8);
    margin-top: 5px;
    border-radius: 5px;
}

</style>
""", unsafe_allow_html=True)

st.markdown("""
<style>

/* ===== SIDEBAR BACKGROUND ===== */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #041c32, #04293a, #064663);
    border-right: 1px solid rgba(0,255,255,0.3);
    box-shadow: 5px 0 25px rgba(0,255,255,0.2);
}

/* ===== SIDEBAR TITLE ===== */
section[data-testid="stSidebar"] h2, 
section[data-testid="stSidebar"] h3 {
    color: #00f5ff !important;
    text-shadow: 0 0 10px #00f5ff,
                 0 0 20px #00f5ff;
    font-weight: 800 !important;
}

/* ===== RADIO CONTAINER ===== */
section[data-testid="stSidebar"] div[role="radiogroup"] label {
    background: rgba(255,255,255,0.05);
    padding: 14px 18px !important;
    border-radius: 16px;
    margin-bottom: 14px;
    font-size: 16px !important;
    font-weight: 600 !important;
    color: #ffffff !important;
    transition: all 0.3s ease;
    border: 1px solid transparent;
    position: relative;
    overflow: hidden;
}

/* Xoá nền trắng bên trong */
section[data-testid="stSidebar"] div[role="radiogroup"] label > div {
    background: transparent !important;
}

/* Hover effect */
section[data-testid="stSidebar"] div[role="radiogroup"] label:hover {
    background: rgba(0,255,255,0.12);
    border: 1px solid #00f5ff;
    box-shadow: 0 0 20px rgba(0,255,255,0.6);
    transform: translateX(6px);
}

/* ===== ACTIVE EFFECT ===== */
section[data-testid="stSidebar"] div[role="radiogroup"] input:checked + div p {
    color: #00f5ff !important;
    font-weight: 700 !important;
    text-shadow: 0 0 10px #00f5ff;
}

/* Thanh neon chạy ngang */
section[data-testid="stSidebar"] div[role="radiogroup"] input:checked + div::before {
    content: "";
    position: absolute;
    left: -100%;
    top: 0;
    width: 100%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(0,255,255,0.6), transparent);
    animation: neonRun 2s linear infinite;
}

@keyframes neonRun {
    from { left: -100%; }
    to { left: 100%; }
}

/* Ẩn chấm radio */
section[data-testid="stSidebar"] div[role="radiogroup"] input {
    display: none;
}

</style>
""", unsafe_allow_html=True)

# =========================
# HEADER
# =========================
st.markdown('<div class="big-title">Di Tích Lịch Sử, Danh Lam Thắng Cảnh Tỉnh Lào Cai </div>', unsafe_allow_html=True)
st.divider()

# =========================
# LOAD DATA
# =========================
try:
    with open("data.json", "r", encoding="utf-8") as f:
        raw_data = json.load(f)
except Exception as e:
    st.error(f"Lỗi đọc file data.json: {e}")
    st.stop()

data = raw_data

# =========================
# MAP ẢNH
# =========================
import os
import streamlit as st

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGE_FOLDER = os.path.join(BASE_DIR, "images")

def load_images_by_prefix(prefix):
    if not os.path.exists(IMAGE_FOLDER):
        st.warning("Không tìm thấy thư mục images")
        return []

    files = os.listdir(IMAGE_FOLDER)

    matched_files = []
    for file in files:
        name = file.lower().strip()
        if name.startswith(prefix.lower()) and name.endswith((".jpg", ".jpeg", ".png")):
            matched_files.append(os.path.join(IMAGE_FOLDER, file))

    return sorted(matched_files)

images = {
    "Đền Thượng": load_images_by_prefix("den_thuong"),
    "Đền Bảo Hà": load_images_by_prefix("den_bao_ha"),
    "Đền Chiềng Ken": load_images_by_prefix("den_chieng_ken"),
    "Đỉnh Fansipan": load_images_by_prefix("fansipan"),
}

# =========================
# SIDEBAR
# =========================

st.sidebar.title("LAO CAI HERITAGE AI")

selected_place = st.sidebar.selectbox(
    "Chọn di tích",
    [p["name"] for p in data]
)

feature = st.sidebar.radio(
    "Tính năng",
    ["Giới thiệu DI TÍCH", "Chatbot AI", "LỊCH TRÌNH DU LỊCH", "THUYẾT MINH DI TÍCH"]
)

place_data = next(p for p in data if p["name"] == selected_place)

# =========================
# GIỚI THIỆU
# =========================
if feature == "Giới thiệu DI TÍCH":

    import streamlit.components.v1 as components
    import base64

    import os

    # mapping tên hiển thị → prefix file ảnh
    mapping = {
        "Den Thuong Lao Cai": "den_thuong",
        "Den Bao Ha": "den_bao_ha",
        "Den Chieng Ken": "den_chieng_ken",
        "Dinh Fansipan": "fansipan"
    }

    prefix = mapping.get(selected_place)

    image_folder = "images"

    image_list = [
        os.path.join(image_folder, f)
        for f in os.listdir(image_folder)
        if prefix and f.startswith(prefix)
    ]

    if image_list:

        images_html = ""

        for img_path in image_list:
            try:
                with open(img_path, "rb") as img_file:
                    encoded = base64.b64encode(img_file.read()).decode()
                images_html += f'<img src="data:image/jpeg;base64,{encoded}" class="slide">'
            except:
                pass

        display_time = 2.5   # mỗi ảnh hiển thị bao nhiêu giây
        total_time = len(image_list) * display_time

        slider_html = f"""
        <style>
        .slider-container {{
            position: relative;
            max-width: 1000px;
            height: 420px;
            margin: 40px auto;
            overflow: hidden;
            border-radius: 25px;
            box-shadow: 0 25px 60px rgba(0,0,0,0.3);
        }}

        .slide {{
            position: absolute;
            width: 100%;
            height: 100%;
            object-fit: cover;
            opacity: 0;
            animation: fadeZoom {total_time}s infinite;
        }}

        @keyframes fadeZoom {{
            0% {{ opacity: 0; transform: scale(1.08); }}
            5% {{ opacity: 1; transform: scale(1); }}
            20% {{ opacity: 1; transform: scale(1); }}
            25% {{ opacity: 0; transform: scale(1.08); }}
            100% {{ opacity: 0; }}
        }}
        """

        # delay tự động theo display_time
        delay = 0
        for i in range(len(image_list)):
            slider_html += f"""
            .slide:nth-child({i+1}) {{
                animation-delay: {delay}s;
            }}
            """
            delay += display_time

        slider_html += f"""
        </style>

        <div class="slider-container">
            {images_html}
        </div>
        """

        components.html(slider_html, height=460)

    else:
        st.warning("⚠ Không có ảnh cho di tích này.")

    # =========================
    # THÔNG TIN
    # =========================
    col1, col2 = st.columns(2)

    with col1:
        st.markdown('<div class="section-title">📖 Mô tả</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="card">{place_data.get("mo_ta","Chưa có dữ liệu.")}</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-title">🏛 Lịch sử</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="card">{place_data.get("lich_su","Chưa có dữ liệu.")}</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="section-title">📌 Địa điểm</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="card">{place_data.get("dia_diem","Chưa có dữ liệu.")}</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-title">🌟 Giá trị văn hóa</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="card">{place_data.get("gia_tri_van_hoa","Chưa có dữ liệu.")}</div>', unsafe_allow_html=True)

# =========================
# TÍNH NĂNG KHÁC
# =========================
elif feature == "Chatbot AI":
    import json
    import streamlit as st
    import time
    import random
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity



    st.markdown("""
    <style>

    /* ===== KHUNG NGOÀI CHAT ===== */
    div[data-testid="stChatInput"] {
        background: rgba(255,255,255,0.55) !important;
        backdrop-filter: blur(15px);
        border-radius: 100px !important;
        padding: 6px 12px !important;
        border: 1px solid rgba(120,180,255,0.4) !important;
        box-shadow: 0 6px 20px rgba(0,120,255,0.15);
        transition: all 0.3s ease;
    }

    /* Hover */
    div[data-testid="stChatInput"]:hover {
        box-shadow: 0 8px 25px rgba(0,120,255,0.25);
    }

    /* Focus */
    div[data-testid="stChatInput"]:focus-within {
        border: 1px solid rgba(0,150,255,0.7) !important;
        box-shadow: 0 0 18px rgba(0,150,255,0.35);
    }

    /* ===== TEXTAREA BÊN TRONG ===== */
    div[data-testid="stChatInput"] textarea {
        background: transparent !important;
        border: none !important;
        outline: none !important;
        box-shadow: none !important;
        border-radius: 40px !important;
        color: #033b5c !important;
        font-size: 17px !important;
    }

    /* Placeholder */
    div[data-testid="stChatInput"] textarea::placeholder {
        color: rgba(0,70,120,0.6) !important;
    }

    /* ===== NÚT GỬI ===== */
    div[data-testid="stChatInput"] button {
        border-radius: 50% !important;
        background: linear-gradient(135deg, #6ec6ff, #2196f3) !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(0,120,255,0.3);
        transition: all 0.3s ease;
    }

    div[data-testid="stChatInput"] button:hover {
        transform: scale(1.08);
        box-shadow: 0 6px 18px rgba(0,120,255,0.4);
    }

    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <style>

    /* XÓA TOÀN BỘ OUTLINE + BORDER MỌI TRẠNG THÁI */
    div[data-testid="stChatInput"] * {
        outline: none !important;
    }

    /* BaseWeb textarea container */
    [data-baseweb="textarea"],
    [data-baseweb="textarea"] > div,
    [data-baseweb="textarea"] textarea {
        border: none !important;
        box-shadow: none !important;
        outline: none !important;
        background: transparent !important;
    }

    /* Khi focus */
    [data-baseweb="textarea"]:focus-within {
        border: none !important;
        box-shadow: none !important;
    }

    /* Textarea focus riêng */
    textarea:focus {
        outline: none !important;
        box-shadow: none !important;
    }

    </style>
    """, unsafe_allow_html=True)
    

    # ==============================
    # LOAD DATA
    # ==============================

    def load_json(file_path="chatai_full.json"):
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)

    data = load_json()

    knowledge_data = [item for item in data if item["di_tich"] != "General"]
    conversation_data = [item for item in data if item["di_tich"] == "General"]

    knowledge_texts = [item["content"] for item in knowledge_data]

    vectorizer = TfidfVectorizer()
    doc_vectors = vectorizer.fit_transform(knowledge_texts)

    # ==============================
    # CHAT HISTORY (GIỮ NGỮ CẢNH)
    # ==============================

    if "messages" not in st.session_state:
        st.session_state.messages = []

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # ==============================
    # XỬ LÝ CHAT
    # ==============================

    if prompt := st.chat_input("Hỏi bất cứ điều gì..."):

        st.session_state.messages.append({"role": "user", "content": prompt})

        with st.chat_message("user"):
            st.markdown(prompt)

        prompt_lower = prompt.lower()

        # ==============================
        # 1️⃣ CHECK GIAO TIẾP
        # ==============================

        reply = None

        for item in conversation_data:
            if any(keyword in prompt_lower for keyword in item["keywords"]):
                reply = item["content"]
                break

        # ==============================
        # 2️⃣ SEARCH RAG
        # ==============================

        if not reply:
            query_vec = vectorizer.transform([prompt])
            similarities = cosine_similarity(query_vec, doc_vectors)
            best_index = similarities.argmax()
            best_score = similarities[0][best_index]

            if best_score > 0.15:
                result = knowledge_data[best_index]
                reply = f"📍 **{result['di_tich']}**\n\n{result['content']}"
            else:
                reply = "Xin lỗi, tôi chỉ cung cấp thông tin về Đền Thượng, Đền Bảo Hà, Đền Chiềng Ken và Đỉnh Fansipan tại Lào Cai."

        # ==============================
        # 3️⃣ HIỆU ỨNG CAO CẤP
        # ==============================

        with st.chat_message("assistant"):

            # --- Hiệu ứng suy nghĩ thông minh ---
            thinking = st.empty()
            dots = ""
            for _ in range(random.randint(2, 4)):
                dots += "."
                thinking.markdown(f"🤖 Đang phân tích{dots}")
                time.sleep(random.uniform(0.3, 0.5))

            thinking.empty()
            time.sleep(random.uniform(0.4, 0.8))

            # --- Streaming mượt ---
            message_placeholder = st.empty()
            full_response = ""
            words = reply.split()

            base_speed = 0.05
            if len(words) > 80:
                base_speed = 0.03
            elif len(words) > 40:
                base_speed = 0.04

            for word in words:
                full_response += word + " "
                message_placeholder.markdown(full_response + "▌")
                time.sleep(random.uniform(base_speed, base_speed + 0.03))

            message_placeholder.markdown(full_response)

        st.session_state.messages.append({"role": "assistant", "content": reply})



elif feature == "LỊCH TRÌNH DU LỊCH":
    st.title("Du Lịch Lào Cai")

    # Đọc dữ liệu
    with open("data.json", "r", encoding="utf-8") as f:
        data = json.load(f)

    # Chọn loại hình
    theme = st.selectbox(
        "Chọn loại du lịch:",
        ["Tâm_Linh", "Check-In"]
    )

    # Lọc địa điểm
    places = [p for p in data if theme in p["type"]]


    total_cost = 0

    for p in places:

        with st.container():
            st.markdown("---")

            # Tiêu đề
            st.markdown(f"## 📍 {p['name']}")

            # Chia 2 cột
            col1, col2 = st.columns([2, 1])

            with col1:
                st.write(f"📌 **Địa điểm:** {p['dia_diem']}")

                # ===== Giờ mở cửa =====
                st.write("🕒 **Giờ mở cửa:**")

                if "opening_hours" in p:

                    if isinstance(p["opening_hours"], dict):
                        with st.expander("Xem chi tiết"):
                            for day, time in p["opening_hours"].items():
                                st.write(f"{day}: {time}")

                    elif isinstance(p["opening_hours"], str):
                        st.write(p["opening_hours"])

                else:
                    st.write("Không có thông tin")

            with col2:
                if isinstance(p["ticket_price"], int):
                    st.markdown(
                        f"""
                        <div style="
                            background-color:#f0f2f6;
                            padding:5px;
                            border-radius:10px;
                            text-align:center;
                        ">
                        <h4>💰 Vé</h4>
                        <h3>{p['ticket_price']:,} VND</h3>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                else:
                    st.write(p["ticket_price"])

                # ===== CHỈ FANSIPAN MỚI CÓ LINK =====
                if "Fansipan" in p["name"]:
                    st.markdown(
                        """
                        <a href="https://vecaptreosapa.com/bang-gia-cap-treo-fansipan/" target="_blank">
                            <button style="
                                margin-top:1px;
                                padding:7px 90px;
                                background:#28a745;
                                color:white;
                                border:none;
                                border-radius:50px;
                                cursor:pointer;
                                font-size:15px;">
                                🔗 Giá vé chi tiết
                            </button>
                        </a>
                        """,
                        unsafe_allow_html=True
                    )

            st.markdown("<br>", unsafe_allow_html=True)


    # 🗺 BẢN ĐỒ
    # ======================

    if places:
        st.markdown("""
        <style>
        .stMap {
            border-radius: 15px;
            overflow: hidden;
        }
        </style>
        """, unsafe_allow_html=True)
        
        df = pd.DataFrame(places)

        if "lat" in df.columns and "lon" in df.columns:

            st.markdown("## 🗺 Our AI Map")

            # Tạo khoảng trắng hai bên để map không full màn hình
            col1, col2, col3 = st.columns([1, 4, 1])

            with col2:
                st.map(
                    df[["lat", "lon"]],
                    height=320,
                    use_container_width=True
                )

        else:
            st.warning("Dữ liệu chưa có tọa độ để hiển thị bản đồ.")


    # ==============================
    # PHẦN CHỈ ĐƯỜNG (dùng full_name)
    # ==============================

    st.markdown("## 📍 Chỉ đường đến địa điểm")

    # Nhập địa chỉ người dùng
    user_address = st.text_input(
        "Chọn điểm bắt đầu..."
    )

    # Chọn điểm đến
    destination_name = st.selectbox(
        "Chọn điểm đến...",
        [p["name"] for p in places]
    )

    # Lấy thông tin địa điểm
    destination = next(p for p in places if p["name"] == destination_name)

    destination_full = destination["full_name"]

    # Chọn phương tiện
    mode_display = st.selectbox(
        "Phương tiện đi lại tốt nhất...",
        ["Ô_Tô", "Đi_Bộ", "Xe_Đạp"]
    )

    mode_map = {
        "Ô_Tô": "driving",
        "Đi_Bộ": "walking",
        "Xe_Đạp": "bicycling"
    }

    travel_mode = mode_map[mode_display]

    import urllib.parse

    if st.button("🚗 Chọn tuyến đường tốt nhất"):
        if user_address:

            origin_encoded = urllib.parse.quote(user_address)
            dest_encoded = urllib.parse.quote(destination_full)

            map_link = (
                "https://www.google.com/maps/dir/?api=1"
                f"&origin={origin_encoded}"
                f"&destination={dest_encoded}"
                f"&travelmode={travel_mode}"
            )

            st.markdown(
                f"👉 **[MỞ GOOGLE MAPS CHỈ ĐƯỜNG]({map_link})**"
            )
        else:
            st.warning("Vui lòng chọn điểm bắt đầu!")




elif feature == "THUYẾT MINH DI TÍCH":

    # ================== CẤU HÌNH TRANG ==================
    st.set_page_config(
        page_title="LAO CAI HERITAGE AI",
        page_icon="logo.png",
        layout="wide"
    )

    # ================== CSS ==================
    st.markdown("""
    <style>
    .main {
        background: linear-gradient(to right, #f8f9fa, #e9ecef);
    }
    .title {
        text-align: center;
        font-size: 42px;
        font-weight: bold;
        color: #2c3e50;
    }
    .subtitle {
        text-align: center;
        font-size: 18px;
        color: gray;
    }
    .stButton>button {
        background-color: #2c3e50;
        color: white;
        border-radius: 8px;
        height: 45px;
        width: 220px;
    }
    </style>
    """, unsafe_allow_html=True)

    # ================== DATA ==================
    di_tich_data = {
        "Đền Thượng Lào Cai": "den_thuong",
        "Đền Bảo Hà": "den_bao_ha",
        "Fansipan": "fansipan",
        "Đền Chiềng Ken": "den_chieng_ken"
    }

    col1, col2 = st.columns([1,2])

    with col1:
        selected = st.selectbox("📍 Chọn di tích", list(di_tich_data.keys()))

        if st.button("🎧 Nghe thuyết minh"):
            doc_path = f"docs/{di_tich_data[selected]}.docx"

            doc = Document(doc_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])

            if full_text.strip():
                tts = gTTS(text=full_text, lang='vi')
                tts.save("temp_audio.mp3")
                audio_file = open("temp_audio.mp3", "rb")
                st.audio(audio_file, format="audio/mp3")
            else:
                st.warning("File không có nội dung!")

    with col2:
        # ===== Load nhiều ảnh theo số đuôi 1-5 =====
        base_name = di_tich_data[selected]
        import time

        image_files = sorted(glob.glob(f"images/{base_name}*.jpg"))

        if image_files:

            slider_html = """
            <style>
            .slider {
                position: relative;
                width: 100%;
                max-width: 900px;
                height: 500px;
                margin: auto;
                overflow: hidden;
                border-radius: 20px;
                box-shadow: 0 10px 30px rgba(0,0,0,0.3);
            }

            .slider img {
                position: absolute;
                width: 100%;
                height: 100%;
                object-fit: cover;
                opacity: 0;
                animation: fade 16s infinite;
            }

            @keyframes fade {
                0% {opacity: 0;}
                5% {opacity: 1;}
                25% {opacity: 1;}
                30% {opacity: 0;}
                100% {opacity: 0;}
            }
            """

            for i in range(len(image_files)):
                delay = i * 4
                slider_html += f"""
                .slider img:nth-child({i+1}) {{
                    animation-delay: {delay}s;
                }}
                """

            slider_html += "</style><div class='slider'>"

            for img in image_files:
                with open(img, "rb") as f:
                    encoded = base64.b64encode(f.read()).decode()
                slider_html += f"<img src='data:image/jpg;base64,{encoded}'>"

            slider_html += "</div>"

            st.markdown(slider_html, unsafe_allow_html=True)

        else:
            st.warning("Chưa tìm thấy ảnh!")

        # ===== Hiển thị nội dung docx =====
        st.markdown("""
        <style>
        .info-card {
            background-color: #ffffff;
            padding: 20px;
            border-radius: 15px;
            box-shadow: 0 4px 12px rgba(0,0,0,0.08);
            margin-bottom: 20px;
        }

        .section-title {
            font-size: 20px;
            font-weight: bold;
            color: #2e7d32;
            margin-bottom: 10px;
        }

        .section-content {
            font-size: 16px;
            line-height: 1.7;
            color: #333;
        }
        </style>
        """, unsafe_allow_html=True)
        doc_path = f"docs/{base_name}.docx"
        doc = Document(doc_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        st.markdown("### 📖 Nội dung thuyết minh")
        st.write(full_text)

    st.divider()
    st.caption("© 2026 - Dự án giới thiệu di tích văn hóa tỉnh Lào Cai")


st.markdown("""
<style>

/* Ghi đè trực tiếp lớp ngoài cùng của chat input */
div[data-testid="stChatInput"] > div {
    border: none !important;
    box-shadow: none !important;
}

/* Khi focus */
div[data-testid="stChatInput"] > div:focus-within {
    border: none !important;
    box-shadow: none !important;
}

/* Ép tất cả con bên trong không được vẽ viền */
div[data-testid="stChatInput"] * {
    box-shadow: none !important;
}

</style>
""", unsafe_allow_html=True)



